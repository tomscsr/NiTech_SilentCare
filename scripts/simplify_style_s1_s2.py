"""
SilentCare - Simplify Style (Sections 1 & 2)
===============================================
Applies exact text replacements to simplify vocabulary in Sections 1 and 2.
Does NOT modify tables, metrics, figure captions, or section headers.

Usage:
    python scripts/simplify_style_s1_s2.py
"""

from pathlib import Path
from docx import Document

SCRIPT_DIR = Path(__file__).resolve().parent
PROJECT_DIR = SCRIPT_DIR.parent
REPORT_PATH = PROJECT_DIR / "results" / "SilentCare_Internship_Report.docx"

# (find, replace) pairs — applied in order
REPLACEMENTS = [
    (
        "felt far-fetched",
        "felt hard to believe",
    ),
    (
        "vocal prosody",
        "tone of voice",
    ),
    (
        "clearly articulated expressions",
        "clear, exaggerated expressions",
    ),
    (
        "micro-expressions, postural tension",
        "small facial movements, tense body posture",
    ),
    (
        "These are hard constraints, not nice-to-haves.",
        "These are real constraints, not optional.",
    ),
    (
        "remains an open and largely unsolved problem",
        "is a problem that research has not solved yet",
    ),
    (
        "superficially similar behavioural patterns",
        "behaviours that look similar on the surface",
    ),
    (
        "more tractable question",
        "simpler question",
    ),
    (
        "was not a retreat from ambition; it was a recognition that",
        "was not giving up on the original goal \u2014 it was accepting that",
    ),
    (
        "clinically meaningful detection of distress states",
        "reliable detection of distress states",
    ),
    (
        "reflecting the clinical reality that caregivers respond similarly to",
        "because caregivers react the same way to",
    ),
    (
        "differentiates itself from existing solutions along five key dimensions",
        "stands apart from existing solutions in five ways",
    ),
    (
        "The unique combination of local privacy-preserving deployment and "
        "non-verbal specialization positions SilentCare as a research prototype "
        "addressing an underserved clinical need, rather than competing directly "
        "with commercial platforms designed for consumer-facing applications.",
        "SilentCare is not trying to compete with commercial tools. It targets "
        "a specific gap: non-verbal individuals in care settings, with no cloud "
        "dependency and no GPU required.",
    ),
]

# Paragraph styles that must NOT be modified
PROTECTED_STYLES = {"Heading 1", "Heading 2", "Heading 3", "Heading 4"}


def apply_replacements(doc):
    """Apply all replacements, operating on runs to preserve formatting."""
    count = 0

    for old_text, new_text in REPLACEMENTS:
        found = False
        for p in doc.paragraphs:
            # Skip protected styles
            sname = p.style.name if p.style else ""
            if sname in PROTECTED_STYLES:
                continue

            # Skip figure captions and table captions
            stripped = p.text.strip()
            if stripped.startswith("Figure ") and " -- " in stripped[:40]:
                continue
            if stripped.startswith("Table ") and " -- " in stripped[:40]:
                continue

            if old_text not in p.text:
                continue

            # Replace within runs (preserves formatting)
            # Collect full text from runs to handle cross-run matches
            full_text = "".join(r.text for r in p.runs)
            if old_text in full_text:
                # Find which runs contain the old text
                # Strategy: rebuild run texts with replacement
                new_full = full_text.replace(old_text, new_text, 1)
                # Distribute new text across existing runs
                _distribute_text(p.runs, new_full)
                found = True
                count += 1
                # Show context around the change
                idx = new_full.find(new_text)
                start = max(0, idx - 20)
                end = min(len(new_full), idx + len(new_text) + 20)
                context = new_full[start:end]
                print(f"  [{count:2d}] ...{context}...")
                break

        if not found:
            print(f"  SKIP: \"{old_text[:60]}...\" (not found or already replaced)")

    return count


def _distribute_text(runs, new_full_text):
    """Distribute new_full_text across existing runs, preserving formatting."""
    if not runs:
        return

    # Put all text in first run, clear the rest
    runs[0].text = new_full_text
    for r in runs[1:]:
        r.text = ""


def main():
    print("=" * 60)
    print("SilentCare - Simplify Style (Sections 1 & 2)")
    print("=" * 60)

    if not REPORT_PATH.exists():
        print(f"ERROR: Report not found at {REPORT_PATH}")
        return

    doc = Document(str(REPORT_PATH))
    print(f"Loaded: {REPORT_PATH}\n")

    count = apply_replacements(doc)

    if count > 0:
        doc.save(str(REPORT_PATH))
        print(f"\n{count} replacements applied.")
        print(f"Saved: {REPORT_PATH}")
    else:
        print("\nNo replacements applied.")

    # Verification pass
    print("\nVerification:")
    doc2 = Document(str(REPORT_PATH))
    for old_text, new_text in REPLACEMENTS:
        old_found = any(old_text in p.text for p in doc2.paragraphs)
        new_found = any(new_text in p.text for p in doc2.paragraphs)
        short = new_text[:55] + "..." if len(new_text) > 55 else new_text
        if new_found and not old_found:
            print(f"  OK   \"{short}\"")
        elif old_found:
            print(f"  FAIL old text still present: \"{old_text[:55]}...\"")
        else:
            print(f"  ???  neither old nor new found")

    print("=" * 60)


if __name__ == "__main__":
    main()

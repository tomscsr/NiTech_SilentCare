"""
SilentCare - Add Section 1.3 (Initial Scope and Why It Changed)
=================================================================
Inserts a new subsection 1.3 into Section 1, between the current
1.2 (Problem Statement) and 1.3 (Internship Objective).

Steps:
  1. Renumber current "1.4 Report Structure" -> "1.5 Report Structure"
  2. Renumber current "1.3 Internship Objective" -> "1.4 Internship Objective"
  3. Insert new Heading 2 + 3 paragraphs before the (now) 1.4

Usage:
    python scripts/add_section_1_3.py
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

SCRIPT_DIR = Path(__file__).resolve().parent
PROJECT_DIR = SCRIPT_DIR.parent
REPORT_PATH = PROJECT_DIR / "results" / "SilentCare_Internship_Report.docx"

BODY_FONT = "Arial"
BODY_SIZE = Pt(11)

# ============================================
# New content
# ============================================

HEADING_TEXT = "1.3 Initial Scope and Why It Changed"

PARA_1 = (
    "The original brief for this internship was broader than what SilentCare "
    "became. The initial objective was to detect not just emotional states but "
    "intentions \u2014 understanding what a non-verbal individual was trying to "
    "communicate, combining facial expressions, body posture, and contextual "
    "cues into a richer behavioural reading. Beyond the technical challenges, "
    "the literature itself suggests that intention detection in non-verbal "
    "individuals remains an open and largely unsolved problem: current models "
    "struggle to reliably distinguish between superficially similar behavioural "
    "patterns, and the absence of large annotated datasets for this population "
    "makes robust training difficult."
)

PARA_2 = (
    "The practical constraints reinforced this conclusion quickly. The multimodal "
    "vision-language models I initially explored were not designed for real-time "
    "inference on consumer hardware. Processing a single frame took several "
    "seconds, and maintaining a continuous video stream was simply not feasible "
    "on CPU. The analysis could produce results, but always too late to be "
    "actionable. A caregiver monitoring a patient cannot wait three seconds per "
    "observation \u2014 by the time the system flags something, the moment has "
    "passed."
)

PARA_3 = (
    "This constraint forced a fundamental redefinition of the problem. Instead "
    "of asking \u2018what is this person trying to express?\u2019, the system "
    "needed to ask \u2018is this person in distress right now?\u2019 \u2014 a "
    "simpler, more tractable question that could be answered within a strict "
    "latency budget. The shift from intention detection to crisis detection was "
    "not a retreat from ambition; it was a recognition that a system that works "
    "in real conditions is more valuable than one that works only on paper. "
    "Everything that followed \u2014 the 4-class schema, the ViT selection, "
    "the fusion design, the alert thresholds \u2014 flows directly from that "
    "initial constraint."
)


def make_body_paragraph(doc, text, anchor_element):
    """Create a body paragraph and insert it before anchor_element."""
    new_p = doc.add_paragraph()
    new_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = new_p.add_run(text)
    run.font.size = BODY_SIZE
    run.font.name = BODY_FONT
    # Move before anchor
    anchor_element.addprevious(new_p._element)
    return new_p


def make_heading2(doc, text, anchor_element):
    """Create a Heading 2 paragraph and insert it before anchor_element."""
    new_p = doc.add_paragraph()
    run = new_p.add_run(text)

    # Apply Heading 2 style via XML (robust against style key issues)
    pPr = new_p._element.get_or_add_pPr()
    pStyle = pPr.get_or_add_pStyle()
    pStyle.set(qn("w:val"), "Heading2")

    # Move before anchor
    anchor_element.addprevious(new_p._element)
    return new_p


def main():
    print("=" * 60)
    print("SilentCare - Add Section 1.3")
    print("=" * 60)

    if not REPORT_PATH.exists():
        print(f"ERROR: Report not found at {REPORT_PATH}")
        return

    doc = Document(str(REPORT_PATH))
    print(f"Loaded: {REPORT_PATH}\n")

    # Check if already inserted
    for p in doc.paragraphs:
        if "1.3" in p.text[:10] and "Initial Scope" in p.text:
            print("  Section 1.3 already exists. No changes needed.")
            return

    # Step 1: Renumber 1.4 -> 1.5
    renamed_14 = False
    for p in doc.paragraphs:
        sname = p.style.name if p.style else ""
        if sname == "Heading 2" and p.text.strip().startswith("1.4 "):
            old = p.text
            if p.runs:
                p.runs[0].text = p.runs[0].text.replace("1.4 ", "1.5 ", 1)
            else:
                p.text = p.text.replace("1.4 ", "1.5 ", 1)
            print(f"  Renumbered: '{old.strip()}' -> '{p.text.strip()}'")
            renamed_14 = True
            break

    if not renamed_14:
        print("  WARNING: '1.4 ...' heading not found for renumbering")

    # Step 2: Renumber 1.3 -> 1.4
    anchor_para = None
    renamed_13 = False
    for p in doc.paragraphs:
        sname = p.style.name if p.style else ""
        if sname == "Heading 2" and p.text.strip().startswith("1.3 "):
            old = p.text
            if p.runs:
                p.runs[0].text = p.runs[0].text.replace("1.3 ", "1.4 ", 1)
            else:
                p.text = p.text.replace("1.3 ", "1.4 ", 1)
            print(f"  Renumbered: '{old.strip()}' -> '{p.text.strip()}'")
            anchor_para = p
            renamed_13 = True
            break

    if not renamed_13 or anchor_para is None:
        print("  ERROR: '1.3 ...' heading not found. Cannot insert new section.")
        return

    # Step 3: Insert new content BEFORE the (now) 1.4 heading
    # Insert in forward order: heading first, para3 last
    # addprevious puts each element before the anchor, so first inserted = topmost
    anchor_el = anchor_para._element

    make_heading2(doc, HEADING_TEXT, anchor_el)
    make_body_paragraph(doc, PARA_1, anchor_el)
    make_body_paragraph(doc, PARA_2, anchor_el)
    make_body_paragraph(doc, PARA_3, anchor_el)

    print(f"\n  Inserted: '{HEADING_TEXT}' + 3 paragraphs")

    doc.save(str(REPORT_PATH))
    print(f"\nSaved: {REPORT_PATH}")

    # Verify
    print("\nVerification:")
    doc2 = Document(str(REPORT_PATH))
    section1_headings = []
    in_s1 = False
    for p in doc2.paragraphs:
        sname = p.style.name if p.style else ""
        if sname == "Heading 1" and p.text.strip().startswith("1."):
            in_s1 = True
            continue
        if sname == "Heading 1" and not p.text.strip().startswith("1."):
            break
        if in_s1 and sname == "Heading 2":
            section1_headings.append(p.text.strip())

    for h in section1_headings:
        ok = True
        if "1.3" in h and "Initial Scope" in h:
            print(f"  OK  {h}")
        elif "1.4" in h and "Internship" in h:
            print(f"  OK  {h}")
        elif "1.5" in h and "Report" in h:
            print(f"  OK  {h}")
        else:
            print(f"  --  {h}")

    print("=" * 60)


if __name__ == "__main__":
    main()

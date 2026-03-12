"""
SilentCare - Simplify Style (Section 5)
=========================================
Applies exact text replacements to simplify vocabulary in Section 5.
Does NOT modify tables, metrics, figure captions, or section headers.

Usage:
    python scripts/simplify_style_s5.py
"""

from pathlib import Path
from docx import Document

SCRIPT_DIR = Path(__file__).resolve().parent
PROJECT_DIR = SCRIPT_DIR.parent
REPORT_PATH = PROJECT_DIR / "results" / "SilentCare_Internship_Report.docx"

REPLACEMENTS = [
    (
        "a pattern jumped out immediately",
        "a pattern stood out immediately",
    ),
    (
        "it simply did not have the capacity to do it",
        "it just could not do it",
    ),
    (
        "which felt like progress until I listened to some of the misclassified samples",
        "which felt like progress until I looked at the mistakes more closely",
    ),
    (
        "acoustically ambiguous",
        "hard to classify by ear",
    ),
    (
        "over-represented. I knew this would lead to overfitting, and it did.",
        "over-represented. I knew this would cause overfitting, and it did.",
    ),
    (
        "processed at key intervals",
        "sampled at regular intervals",
    ),
    (
        "creating labeled training data for future incremental fine-tuning",
        "creating labeled data for future retraining",
    ),
]

PROTECTED_STYLES = {"Heading 1", "Heading 2", "Heading 3", "Heading 4"}


def apply_replacements(doc):
    count = 0
    for old_text, new_text in REPLACEMENTS:
        found = False
        for p in doc.paragraphs:
            sname = p.style.name if p.style else ""
            if sname in PROTECTED_STYLES:
                continue
            stripped = p.text.strip()
            if stripped.startswith("Figure ") and " -- " in stripped[:40]:
                continue
            if stripped.startswith("Table ") and " -- " in stripped[:40]:
                continue
            if old_text not in p.text:
                continue

            full_text = "".join(r.text for r in p.runs)
            if old_text in full_text:
                new_full = full_text.replace(old_text, new_text, 1)
                p.runs[0].text = new_full
                for r in p.runs[1:]:
                    r.text = ""
                found = True
                count += 1
                idx = new_full.find(new_text)
                start = max(0, idx - 20)
                end = min(len(new_full), idx + len(new_text) + 20)
                print(f"  [{count:2d}] ...{new_full[start:end]}...")
                break

        if not found:
            print(f"  SKIP: \"{old_text[:60]}\" (not found)")

    return count


def main():
    print("=" * 60)
    print("SilentCare - Simplify Style (Section 5)")
    print("=" * 60)

    if not REPORT_PATH.exists():
        print(f"ERROR: Report not found at {REPORT_PATH}")
        return

    doc = Document(str(REPORT_PATH))
    print(f"Loaded: {REPORT_PATH}\n")

    count = apply_replacements(doc)

    if count > 0:
        doc.save(str(REPORT_PATH))
        print(f"\n{count} replacements applied.")
        print(f"Saved: {REPORT_PATH}")
    else:
        print("\nNo replacements applied.")

    print("\nVerification:")
    doc2 = Document(str(REPORT_PATH))
    for old_text, new_text in REPLACEMENTS:
        old_found = any(old_text in p.text for p in doc2.paragraphs)
        new_found = any(new_text in p.text for p in doc2.paragraphs)
        short = new_text[:55] + "..." if len(new_text) > 55 else new_text
        if new_found and not old_found:
            print(f"  OK   \"{short}\"")
        elif old_found:
            print(f"  FAIL old still present: \"{old_text[:55]}\"")
        else:
            print(f"  ???  neither old nor new found")

    print("=" * 60)


if __name__ == "__main__":
    main()

"""
SilentCare - Simplify Style (Section 2 fix + Section 3)
=========================================================
Applies exact text replacements to simplify vocabulary.
Does NOT modify tables, metrics, figure captions, or section headers.

Usage:
    python scripts/simplify_style_s3.py
"""

from pathlib import Path
from docx import Document

SCRIPT_DIR = Path(__file__).resolve().parent
PROJECT_DIR = SCRIPT_DIR.parent
REPORT_PATH = PROJECT_DIR / "results" / "SilentCare_Internship_Report.docx"

REPLACEMENTS = [
    (
        "expect exaggerated, clear, exaggerated expressions",
        "expect clear, exaggerated expressions",
    ),
    (
        "a fundamental reframing of the problem was proposed: the most impactful "
        "use case is not continuous fine-grained emotion labeling, but reliable "
        "distress detection for people who cannot communicate verbally.",
        "I decided to reframe the problem entirely: the most useful thing this "
        "system could do was not label emotions continuously, but reliably detect "
        "distress in people who cannot communicate verbally.",
    ),
    (
        "This pivot entailed two major changes",
        "This meant two major changes",
    ),
    (
        "This mapping reflects the clinical reality that a caregiver\u2019s "
        "response to fear and sadness is identical (investigate and comfort)",
        "This mapping reflects the fact that a caregiver reacts the same way "
        "to fear and sadness (check and comfort the patient)",
    ),
    (
        "Four quantitative objectives were established to guide development",
        "I set four concrete objectives to guide the project",
    ),
    (
        "spurious notifications through temporal smoothing and confidence gating",
        "unnecessary alerts through timing filters and confidence checks",
    ),
]

PROTECTED_STYLES = {"Heading 1", "Heading 2", "Heading 3", "Heading 4"}


def apply_replacements(doc):
    count = 0
    for old_text, new_text in REPLACEMENTS:
        found = False
        for p in doc.paragraphs:
            sname = p.style.name if p.style else ""
            if sname in PROTECTED_STYLES:
                continue
            stripped = p.text.strip()
            if stripped.startswith("Figure ") and " -- " in stripped[:40]:
                continue
            if stripped.startswith("Table ") and " -- " in stripped[:40]:
                continue

            if old_text not in p.text:
                continue

            full_text = "".join(r.text for r in p.runs)
            if old_text in full_text:
                new_full = full_text.replace(old_text, new_text, 1)
                p.runs[0].text = new_full
                for r in p.runs[1:]:
                    r.text = ""
                found = True
                count += 1
                idx = new_full.find(new_text)
                start = max(0, idx - 20)
                end = min(len(new_full), idx + len(new_text) + 20)
                print(f"  [{count:2d}] ...{new_full[start:end]}...")
                break

        if not found:
            # Try with straight apostrophe as fallback
            alt_old = old_text.replace("\u2019", "'")
            for p in doc.paragraphs:
                sname = p.style.name if p.style else ""
                if sname in PROTECTED_STYLES:
                    continue
                if alt_old not in p.text:
                    continue
                full_text = "".join(r.text for r in p.runs)
                if alt_old in full_text:
                    new_full = full_text.replace(alt_old, new_text, 1)
                    p.runs[0].text = new_full
                    for r in p.runs[1:]:
                        r.text = ""
                    found = True
                    count += 1
                    idx = new_full.find(new_text)
                    start = max(0, idx - 20)
                    end = min(len(new_full), idx + len(new_text) + 20)
                    print(f"  [{count:2d}] ...{new_full[start:end]}...")
                    break

        if not found:
            print(f"  SKIP: \"{old_text[:60]}\" (not found)")

    return count


def main():
    print("=" * 60)
    print("SilentCare - Simplify Style (Section 2 fix + Section 3)")
    print("=" * 60)

    if not REPORT_PATH.exists():
        print(f"ERROR: Report not found at {REPORT_PATH}")
        return

    doc = Document(str(REPORT_PATH))
    print(f"Loaded: {REPORT_PATH}\n")

    count = apply_replacements(doc)

    if count > 0:
        doc.save(str(REPORT_PATH))
        print(f"\n{count} replacements applied.")
        print(f"Saved: {REPORT_PATH}")
    else:
        print("\nNo replacements applied.")

    # Verification
    print("\nVerification:")
    doc2 = Document(str(REPORT_PATH))
    for old_text, new_text in REPLACEMENTS:
        old_found = any(old_text in p.text for p in doc2.paragraphs)
        new_found = any(new_text in p.text for p in doc2.paragraphs)
        short = new_text[:55] + "..." if len(new_text) > 55 else new_text
        if new_found and not old_found:
            print(f"  OK   \"{short}\"")
        elif old_found:
            print(f"  FAIL old still present: \"{old_text[:55]}\"")
        else:
            print(f"  ???  neither old nor new found")

    print("=" * 60)


if __name__ == "__main__":
    main()

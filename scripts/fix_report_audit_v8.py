"""
SilentCare - Fix Report Audit v8
=================================
Fixes 5 issues in SilentCare_Internship_Report.docx:
  1. Table 10.1: Remove corrupted 4th/5th columns
  2. Appendix A: Verify EfficientNet-B2/MobileNetV3 columns exist
  3. Table 8.2c: Correct ViT trpakov inference 114.3 -> 117.6 ms
  4. Table 8.2c: Correct ResNet50 inference 45.1 -> 40.4 ms
  5. Section 4.2.3: Correct sampling description + Table 8.2c caption

Usage:
    python scripts/fix_report_audit_v8.py
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

SCRIPT_DIR = Path(__file__).resolve().parent
PROJECT_DIR = SCRIPT_DIR.parent
REPORT_PATH = PROJECT_DIR / "results" / "SilentCare_Internship_Report.docx"

BODY_FONT = "Arial"
BODY_SIZE = Pt(11)
TABLE_FONT_SIZE = Pt(10)
CAPTION_SIZE = Pt(10)
HEADER_BG = "D5E8F0"


def make_run(paragraph, text, bold=None, italic=None, size=None, font_name=None):
    run = paragraph.add_run(text)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if size is not None:
        run.font.size = size
    if font_name is not None:
        run.font.name = font_name
    return run


def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def format_header_cell(cell, text):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    make_run(p, text, bold=True, size=TABLE_FONT_SIZE, font_name=BODY_FONT)
    set_cell_shading(cell, HEADER_BG)


def format_body_cell(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    make_run(p, text, bold=bold, size=TABLE_FONT_SIZE, font_name=BODY_FONT)


# ============================================
# FIX 1: Table 10.1 - Remove corrupted columns
# ============================================

def fix_table_10_1(doc):
    """Remove spurious 4th+ columns from Comparison with Objectives table."""
    table = None
    for t in doc.tables:
        if (t.rows[0].cells[0].text.strip() == "Objective" and
                t.rows[0].cells[1].text.strip() == "Target"):
            table = t
            break

    if table is None:
        print("  WARNING: Table 10.1 not found")
        return False

    ncols = len(table.rows[0].cells)
    if ncols <= 3:
        print("  Table 10.1: already 3 columns, nothing to fix")
        return False

    # Remove all columns beyond the 3rd
    for row in table.rows:
        tcs = row._tr.findall(qn("w:tc"))
        while len(tcs) > 3:
            row._tr.remove(tcs[-1])
            tcs = row._tr.findall(qn("w:tc"))

    # Verify
    new_ncols = len(table.rows[0].cells)
    print(f"  FIX 1: Table 10.1 trimmed from {ncols} to {new_ncols} columns")
    return True


# ============================================
# FIX 2: Appendix A - Verify/add new columns
# ============================================

def fix_appendix_a(doc):
    """Ensure Table A.1 has EfficientNet-B2 and MobileNetV3 columns."""
    table = None
    for t in doc.tables:
        if (len(t.rows) >= 10 and
                "Parameter" in t.rows[0].cells[0].text and
                "Audio" in t.rows[0].cells[1].text):
            table = t
            break

    if table is None:
        print("  WARNING: Table A.1 not found")
        return False

    ncols = len(table.rows[0].cells)

    # First, strip any corrupted extra columns (beyond 3 base + 2 new = 5)
    if ncols > 5:
        for row in table.rows:
            tcs = row._tr.findall(qn("w:tc"))
            while len(tcs) > 3:
                row._tr.remove(tcs[-1])
                tcs = row._tr.findall(qn("w:tc"))
        ncols = 3

    if ncols == 5:
        # Check if they already have correct content
        h3 = table.rows[0].cells[3].text.strip()
        h4 = table.rows[0].cells[4].text.strip()
        if "EfficientNet" in h3 and "MobileNet" in h4:
            print("  FIX 2: Table A.1 already has correct 5 columns")
            return False
        else:
            # Remove the bad columns, re-add
            for row in table.rows:
                tcs = row._tr.findall(qn("w:tc"))
                while len(tcs) > 3:
                    row._tr.remove(tcs[-1])
                    tcs = row._tr.findall(qn("w:tc"))
            ncols = 3

    if ncols == 3:
        # Add 2 new columns
        effnet_data = [
            "EfficientNet-B2*",
            "EfficientNet-B2 (fine-tuned) + Head*",
            "PyTorch",
            "224x224x3 (RGB)",
            "4 (softmax via logits)",
            "30",
            "(early stopped)",
            "32",
            "0.0001 (head) / 0.00001 (backbone)",
            "AdamW (weight_decay=1e-4)",
            "ReduceLROnPlateau (factor=0.5)",
            "8",
            "0.5 (projection)",
            "No (EfficientNet has built-in BN)",
            "HFlip, rotation, color jitter",
            "WeightedRandomSampler",
            "RAF-DB native train/test",
            "RAF-DB test partition",
        ]

        mobilenet_data = [
            "MobileNetV3-Large*",
            "MobileNetV3-Large (fine-tuned) + Head*",
            "PyTorch",
            "224x224x3 (RGB)",
            "4 (softmax via logits)",
            "30",
            "(early stopped)",
            "32",
            "0.0001 (head) / 0.00001 (backbone)",
            "AdamW (weight_decay=1e-4)",
            "ReduceLROnPlateau (factor=0.5)",
            "8",
            "0.5 (projection)",
            "No (MobileNetV3 has built-in BN)",
            "HFlip, rotation, color jitter",
            "WeightedRandomSampler",
            "RAF-DB native train/test",
            "RAF-DB test partition",
        ]

        for row_idx, row in enumerate(table.rows):
            for col_data in [effnet_data, mobilenet_data]:
                row._tr.add_tc()
                if row_idx == 0:
                    format_header_cell(row.cells[-1], col_data[row_idx])
                else:
                    format_body_cell(row.cells[-1], col_data[row_idx])

        print(f"  FIX 2: Table A.1 expanded from 3 to 5 columns")

    # Update caption
    for p in doc.paragraphs:
        if "Table A.1 --" in p.text:
            p.clear()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            make_run(p,
                     "Table A.1 -- Training hyperparameters. *ResNet50, EfficientNet-B2, "
                     "and MobileNetV3 were trained for comparison; production uses ViT HuggingFace "
                     "(trpakov/vit-face-expression, pre-trained, no local training).",
                     italic=True, size=CAPTION_SIZE, font_name=BODY_FONT)
            break

    return True


# ============================================
# FIX 3 & 4: Table 8.2c inference corrections
# ============================================

def fix_table_8_2c_inference(doc):
    """Correct ViT trpakov (114.3->117.6) and ResNet50 (45.1->40.4) in Table 8.2c."""
    table = None
    for t in doc.tables:
        if (len(t.rows) >= 6 and
                len(t.rows[0].cells) >= 7 and
                "Model" in t.rows[0].cells[0].text and
                "Rank" in t.rows[0].cells[6].text):
            table = t
            break

    if table is None:
        print("  WARNING: Table 8.2c not found")
        return False

    fixed = False
    for row in table.rows[1:]:
        model_name = row.cells[0].text.strip()
        inference_cell = row.cells[5]
        current_val = inference_cell.text.strip()

        if "trpakov" in model_name and current_val == "114.3":
            format_body_cell(inference_cell, "117.6")
            print(f"  FIX 3: ViT trpakov inference: 114.3 -> 117.6 ms")
            fixed = True

        if "ResNet50" in model_name and current_val == "45.1":
            format_body_cell(inference_cell, "40.4")
            print(f"  FIX 4: ResNet50 inference: 45.1 -> 40.4 ms")
            fixed = True

    if not fixed:
        print("  FIX 3/4: Inference values may already be correct or not found")

    return fixed


# ============================================
# FIX 5: Section 4.2.3 sampling description
# ============================================

def fix_section_4_2_3(doc):
    """Update sampling description in Section 4.2.3 and Table 8.2c caption."""
    fixed = False

    old_sampling = (
        "After per-source balanced sampling (equal representation per source per class), "
        "the benchmark comprises 9,204 images (3,068 per source)."
    )
    new_sampling = (
        "After balanced sampling ensuring equal representation per source (3,068 images "
        "per source, 3 sources), the benchmark comprises 9,204 images. Note that class "
        "distribution within the benchmark reflects the natural class proportions of the "
        "source datasets, as required for realistic cross-domain generalization evaluation."
    )

    for p in doc.paragraphs:
        if old_sampling in p.text:
            new_text = p.text.replace(old_sampling, new_sampling)
            p.clear()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            make_run(p, new_text, size=BODY_SIZE, font_name=BODY_FONT)
            print("  FIX 5a: Section 4.2.3 sampling description updated")
            fixed = True
            break

    # Update Table 8.2c caption
    old_caption_fragment = "Balanced sampling ensures equal representation per source."
    new_caption_fragment = "Balanced sampling ensures equal representation per source (3,068 images per source)."

    for p in doc.paragraphs:
        if "Table 8.2c --" in p.text and old_caption_fragment in p.text:
            new_text = p.text.replace(old_caption_fragment, new_caption_fragment)
            p.clear()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            make_run(p, new_text, italic=True, size=CAPTION_SIZE, font_name=BODY_FONT)
            print("  FIX 5b: Table 8.2c caption updated")
            fixed = True
            break

    if not fixed:
        print("  FIX 5: Text not found (may already be fixed)")

    return fixed


# ============================================
# Main
# ============================================

def main():
    print("=" * 60)
    print("SilentCare - Fix Report Audit v8")
    print("=" * 60)

    if not REPORT_PATH.exists():
        print(f"ERROR: Report not found at {REPORT_PATH}")
        return

    doc = Document(str(REPORT_PATH))
    print(f"Loaded: {REPORT_PATH}\n")

    fix_table_10_1(doc)
    fix_appendix_a(doc)
    fix_table_8_2c_inference(doc)
    fix_section_4_2_3(doc)

    doc.save(str(REPORT_PATH))
    print(f"\nSaved: {REPORT_PATH}")

    # Verify
    print("\nVerification:")
    doc2 = Document(str(REPORT_PATH))
    for t in doc2.tables:
        if t.rows[0].cells[0].text.strip() == "Objective":
            print(f"  Table 10.1: {len(t.rows[0].cells)} cols "
                  f"({'OK' if len(t.rows[0].cells) == 3 else 'FAIL'})")
        if ("Parameter" in t.rows[0].cells[0].text and
                "Audio" in t.rows[0].cells[1].text):
            print(f"  Table A.1: {len(t.rows[0].cells)} cols "
                  f"({'OK' if len(t.rows[0].cells) >= 5 else 'FAIL'})")
        if ("Model" in t.rows[0].cells[0].text and
                len(t.rows[0].cells) >= 7 and
                "Rank" in t.rows[0].cells[6].text):
            for row in t.rows[1:]:
                name = row.cells[0].text.strip()
                inf = row.cells[5].text.strip()
                if "trpakov" in name:
                    print(f"  Table 8.2c ViT trpakov inference: {inf} "
                          f"({'OK' if inf == '117.6' else 'FAIL'})")
                if "ResNet50" in name:
                    print(f"  Table 8.2c ResNet50 inference: {inf} "
                          f"({'OK' if inf == '40.4' else 'FAIL'})")

    print("=" * 60)


if __name__ == "__main__":
    main()
